import os
from typing import Optional
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook


class DocumentParser:
    """Service for parsing different document types"""

    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """Parse PDF file and extract text"""
        try:
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
            return text.strip()
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")

    @staticmethod
    def parse_docx(file_path: str) -> str:
        """Parse DOCX file and extract text"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text += row_text + "\n"

            return text.strip()
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")

    @staticmethod
    def parse_xlsx(file_path: str) -> str:
        """Parse Excel file and extract text from all sheets"""
        try:
            workbook = load_workbook(file_path, data_only=True)
            text = ""

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"=== Sheet: {sheet_name} ===\n"

                for row in sheet.iter_rows(values_only=True):
                    # Filter out None values and convert to strings
                    row_values = [str(cell) for cell in row if cell is not None]
                    if row_values:
                        text += " | ".join(row_values) + "\n"

                text += "\n"

            return text.strip()
        except Exception as e:
            raise ValueError(f"Failed to parse Excel: {str(e)}")

    @staticmethod
    def parse_txt(file_path: str) -> str:
        """Parse plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read().strip()
        except Exception as e:
            raise ValueError(f"Failed to parse text file: {str(e)}")

    def parse_document(self, file_path: str, file_type: str) -> str:
        """Parse document based on file type"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_type = file_type.lower()

        if file_type == "pdf":
            return self.parse_pdf(file_path)
        elif file_type in ["docx", "doc"]:
            return self.parse_docx(file_path)
        elif file_type in ["xlsx", "xls"]:
            return self.parse_xlsx(file_path)
        elif file_type == "txt":
            return self.parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
